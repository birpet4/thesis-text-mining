import spacy
import spacy
import pandas as pd
import numpy as np
from spacy.util import minibatch, compounding
import json
import random
from pandas import DataFrame
from spacy import displacy
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX


def main(doc, document):
    # output_path = Path("test/if30.svg")
    # svg = displacy.render(doc, style='dep')
    # with output_path.open("w", encoding="utf-8") as fh:
    #     fh.write(svg)

    mark = ""
    advcl = ""
    position = 0
    tok = ""
    punct = 0
    lastp = 0
    lastlast = 0
    if_pos = 0
    for token in doc:
        if token.text == "If":
            if_pos = token.idx
            mark = token.head.text

        if token.dep_ == "punct":
            punct = token.idx
            if lastp != 0:
                lastlast = token.idx

        if token.text == mark:
            if token.dep_ == "advcl" and "If" in [children.text for children in token.children]:
                advcl = token.head.text
            elif token.dep_ == "ROOT" and mark in [children.text for children in token.children]:
                position = token.idx

        if token.text == advcl:
            if mark in [children.text for children in token.children] and token.dep_ == "ROOT":
                position = token.idx
                tok = token.text
                lastp = punct

            elif advcl in [children.text for children in token.children]:
                advcl = token.head.text

    # print(lastp, tok, position)
    # print("If the Collateral Manager or the Issuer were to fail to, or be unable to, obtain such authorisation, the Collateral Manager may not be able to continue to manage the Issuer’s assets, or its ability to do so may be impaired. Any regulatory changes arising from implementation of the AIFMD (or otherwise) that impairs the ability of the Collateral Manager to manage the Issuer’s assets may adversely affect the Issuer’s ability to carry out its investment strategy and achieve its investment objective."[
    #       lastp:])
    # print("If the Collateral Manager or the Issuer were to fail to, or be unable to, obtain such authorisation, the Collateral Manager may not be able to continue to manage the Issuer’s assets, or its ability to do so may be impaired. Any regulatory changes arising from implementation of the AIFMD (or otherwise) that impairs the ability of the Collateral Manager to manage the Issuer’s assets may adversely affect the Issuer’s ability to carry out its investment strategy and achieve its investment objective."[
    #       :lastp])
    # with open('fold/if_statements.txt', 'r', encoding='utf-8') as file:
    #     Lines = file.readlines()
    #     i = 0
    #     for line in Lines:
    #         doc = nlp(line)
    #         output_path = Path("test/if" + str(i) + ".svg")
    #         svg = displacy.render(doc, style='dep')
    #         with output_path.open("w", encoding="utf-8") as fh:
    #             fh.write(svg)
    #         i = i + 1
    if lastp == 0:
        document.add_paragraph(doc.text)
    else:
        p2 = document.add_paragraph()
        before = doc.text[1:if_pos]
        p2.add_run(before)
        color1 = doc.text[if_pos:lastp]
        font = p2.add_run(color1).font
        font.highlight_color = WD_COLOR_INDEX.YELLOW

        color2 = doc.text[lastp:lastlast]
        font = p2.add_run(color2).font
        font.highlight_color = WD_COLOR_INDEX.GREEN


if __name__ == "__main__":
    nlp = spacy.load("en_core_web_sm")
    document = Document()
    with open('sample/1acf000d.p.html.9e269fc7.txt', 'r', encoding='utf-8') as file:
        Lines = file.readlines()
        for line in Lines:
            doc = nlp(line)
            main(doc, document)

    document.save('t1.docx')
